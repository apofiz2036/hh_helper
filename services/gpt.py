import requests
from docx import Document
import os

import logging
from utils.logging_config import get_logger

from config import YANDEX_API_KEY, YANDEX_FOLDER_ID


logger = get_logger(__name__)

def load_prompt():
    """
    Загружает промпт из файла prompt.txt
    """
    try:
        file_name = 'texts/prompt.txt'
        with open(file_name, 'r', encoding='utf-8') as file:
            return file.read()
        
    except Exception as e:
        error_message = f"Ошибка при загрузке промпта {e}"
        print(error_message)
        logger.error(error_message)


def load_resume():
    """
    Загружает описание вакансии из resume.docx
    """
    try:
        doc = Document('texts/resume.docx')
        text = ''

        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    
        return text
    
    except Exception as e:
        error_message = f"Ошибка при загрузке резюме {e}"
        print(error_message)
        logger.error(error_message)


def load_preferences():
    """
    Загружает описание вакансии из preferences.docx
    """
    try:
        doc = Document('texts/preferences.docx')
        text = ''

        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    
        return text
    
    except Exception as e:
        error_message = f"Ошибка при загрузке предпочтений {e}"
        print(error_message)
        logger.error(error_message)


def ask_gpt(vacancy_data):
    """
    Отправляет данные  в YandexGPT для анализа
    Возвращает текстовый ответ с оценкой вакансии
    """
    url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Api-Key {YANDEX_API_KEY}",
        "x-folder-id": YANDEX_FOLDER_ID,
    }

    # Загружаем данные резюме
    prompt = load_prompt()
    resume = load_resume()
    preferences = load_preferences()

    prompt_text = f"""
        {prompt}

        Описание вакансии:
        {vacancy_data}

        Моё резюме:
        {resume}

        Мои предпочтения:
        {preferences}
    """ 

    data = {
        "modelUri": f"gpt://{YANDEX_FOLDER_ID}/yandexgpt-lite",
        "messages": [
            {
                "role": "system", 
                "text": "Ты опытный карьерный консультант, твоя задача оценить на сколько данная ваканся подходит этому пользователю"
            },
            {"role": "user", "text": prompt_text},
        ]
    }

    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()

        json_data = response.json()
        result_text = json_data["result"]["alternatives"][0]["message"]["text"]

        return result_text
    
    except Exception as e:
        error_message = f"Ошибка при запросе к YandexGPT: {e}"
        print(error_message)
        logger.warning(error_message)
        return None
